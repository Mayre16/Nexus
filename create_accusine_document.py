#!/usr/bin/env python3
"""
Script para crear documento Word con información técnica del AccuSine PCS+ 200A IP54
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, fill_color):
    """Aplica color de fondo a una celda"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_color)
    tcPr.append(shading)

def create_document():
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # =====================================================
    # PORTADA
    # =====================================================
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title_run = title.add_run("INFORMACIÓN TÉCNICA")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("PowerLogic AccuSine PCS+")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(24)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle2 = doc.add_paragraph()
    subtitle2_run = subtitle2.add_run("Filtro Activo de Armónicos")
    subtitle2_run.font.size = Pt(18)
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    models = doc.add_paragraph()
    models_run = models.add_run("Modelo a Cotizar:")
    models_run.bold = True
    models_run.font.size = Pt(16)
    models.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    model1 = doc.add_paragraph()
    model1_run = model1.add_run("PCSP200D5IP54")
    model1_run.bold = True
    model1_run.font.size = Pt(20)
    model1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    model2 = doc.add_paragraph()
    model2_run = model2.add_run("200A - 380/480V - IP54")
    model2_run.font.size = Pt(14)
    model2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    brand = doc.add_paragraph()
    brand_run = brand.add_run("Schneider Electric")
    brand_run.bold = True
    brand_run.font.size = Pt(16)
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # =====================================================
    # ÍNDICE
    # =====================================================
    doc.add_heading("ÍNDICE", level=1)
    
    toc_items = [
        "1. Introducción al Sistema AccuSine PCS+",
        "2. Especificaciones Técnicas del PCSP200D5IP54",
        "3. Generación de Calor y Disipación Térmica",
        "4. Especificaciones de Espacio e Instalación",
        "5. Características y Mejoras al Sistema",
        "6. Normativas y Certificaciones"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
    
    doc.add_page_break()
    
    # =====================================================
    # SECCIÓN 1: INTRODUCCIÓN
    # =====================================================
    doc.add_heading("1. Introducción al Sistema AccuSine PCS+", level=1)
    
    intro_text = """El PowerLogic AccuSine PCS+ es un filtro activo de armónicos (AHF) diseñado para estabilizar redes eléctricas industriales mediante la inyección dinámica de corriente en tiempo real. Este sistema proporciona:

• Mitigación de armónicos
• Corrección del factor de potencia
• Balanceo de carga

El AccuSine PCS+ inyecta corriente armónica complementaria para cancelar los armónicos en el sistema de distribución eléctrica, lo que resulta en una mayor confiabilidad de la red eléctrica y reducción de costos operativos."""
    
    doc.add_paragraph(intro_text)
    
    doc.add_heading("Principio de Operación", level=2)
    
    operation_text = """El filtro activo AccuSine PCS+ funciona inyectando corriente armónica opuesta en el lado de la fuente de la carga. Esto cancela efectivamente las corrientes armónicas generadas por cargas no lineales, como variadores de frecuencia, UPS, equipos de soldadura y otros dispositivos electrónicos de potencia.

El sistema monitorea continuamente la corriente de línea mediante transformadores de corriente (CTs) y calcula en tiempo real la compensación necesaria para cada armónico individual hasta el orden 51."""
    
    doc.add_paragraph(operation_text)
    
    doc.add_page_break()
    
    # =====================================================
    # SECCIÓN 2: ESPECIFICACIONES TÉCNICAS
    # =====================================================
    doc.add_heading("2. Especificaciones Técnicas del PCSP200D5IP54", level=1)
    
    doc.add_heading("Descripción General", level=2)
    
    pcs200_desc = """El modelo PCSP200D5IP54 es un filtro activo de armónicos con una corriente nominal de salida de 200 Amperios. La designación IP54 indica protección contra polvo limitada y protección contra salpicaduras de agua desde cualquier dirección. Es un gabinete de piso diseñado para aplicaciones industriales pesadas y entornos de misión crítica.

Este hardware es una solución de corrección de energía activa de la familia AccuSine, diseñado específicamente para condiciones eléctricas severas en aplicaciones industriales pesadas."""
    
    doc.add_paragraph(pcs200_desc)
    
    doc.add_heading("Especificaciones Eléctricas", level=2)
    
    # Tabla de especificaciones eléctricas
    table_elec = doc.add_table(rows=13, cols=2)
    table_elec.style = 'Table Grid'
    
    specs_elec = [
        ("Parámetro", "Especificación"),
        ("Referencia del Producto", "PCSP200D5IP54"),
        ("Corriente Nominal de Salida (RMS)", "200 A"),
        ("Potencia Reactiva", "166 kVAR @ 480V"),
        ("Rango de Voltaje de Red", "380-480 V AC"),
        ("Frecuencia de Red", "50/60 Hz ±3 Hz (autodetección)"),
        ("Conexión", "Trifásica 3 hilos o 4 hilos"),
        ("THDi de Salida", "≤ 3%"),
        ("Corrección de Armónicos", "Hasta el 51° armónico"),
        ("Tiempo de Respuesta (Armónicos)", "2 ciclos"),
        ("Tiempo de Respuesta (Factor de Potencia)", "1/4 ciclo"),
        ("Factor de Potencia Objetivo", "Configurable (leading o lagging)"),
        ("Capacidad de Paralelismo", "Hasta 10 unidades (máx. 3000A)")
    ]
    
    for i, (param, spec) in enumerate(specs_elec):
        row = table_elec.rows[i]
        row.cells[0].text = param
        row.cells[1].text = spec
        if i == 0:
            set_cell_shading(row.cells[0], "003366")
            set_cell_shading(row.cells[1], "003366")
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    doc.add_heading("Especificaciones Físicas", level=2)
    
    # Tabla de especificaciones físicas
    table_phys = doc.add_table(rows=9, cols=2)
    table_phys.style = 'Table Grid'
    
    specs_phys = [
        ("Parámetro", "Especificación"),
        ("Grado de Protección", "IP54"),
        ("Tipo de Montaje", "Piso (floor standing)"),
        ("Entrada de Cables", "Superior o Inferior"),
        ("Alto", "2100 mm (82.7 in)"),
        ("Ancho", "900 mm (35.4 in)"),
        ("Profundidad", "600 mm (23.6 in)"),
        ("Peso", "402 kg (886 lbs)"),
        ("Color del Gabinete", "RAL7035 Gris Claro")
    ]
    
    for i, (param, spec) in enumerate(specs_phys):
        row = table_phys.rows[i]
        row.cells[0].text = param
        row.cells[1].text = spec
        if i == 0:
            set_cell_shading(row.cells[0], "003366")
            set_cell_shading(row.cells[1], "003366")
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    doc.add_heading("Conexiones Eléctricas", level=2)
    
    # Tabla de conexiones
    table_conn = doc.add_table(rows=7, cols=2)
    table_conn.style = 'Table Grid'
    
    specs_conn = [
        ("Conexión", "Especificación"),
        ("Terminales de Potencia", "Espárragos M12"),
        ("Terminales de Tierra", "Espárragos M8"),
        ("Separación entre Espárragos", "44.5 mm (1.75 in) centro a centro"),
        ("Temperatura de Cable Permitida", "75°C, 90°C"),
        ("Par de Apriete (Potencia)", "37.0 N·m (327.5 lb-in)"),
        ("Par de Apriete (Tierra)", "18.2 N·m (161.1 lb-in)")
    ]
    
    for i, (param, spec) in enumerate(specs_conn):
        row = table_conn.rows[i]
        row.cells[0].text = param
        row.cells[1].text = spec
        if i == 0:
            set_cell_shading(row.cells[0], "003366")
            set_cell_shading(row.cells[1], "003366")
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # =====================================================
    # SECCIÓN 3: GENERACIÓN DE CALOR
    # =====================================================
    doc.add_heading("3. Generación de Calor y Disipación Térmica", level=1)
    
    heat_intro = """El filtro activo genera calor significativo durante su operación. Es fundamental asegurar una ventilación adecuada en el área de instalación para mantener las condiciones operativas óptimas."""
    
    doc.add_paragraph(heat_intro)
    
    doc.add_heading("Carga Térmica del PCSP200D5IP54", level=2)
    
    # Tabla de calor
    heat_table = doc.add_table(rows=5, cols=2)
    heat_table.style = 'Table Grid'
    
    heat_data = [
        ("Parámetro", "Valor"),
        ("Carga Térmica @ 380V", "3.9 kW"),
        ("Carga Térmica @ 480V", "4.3 kW"),
        ("Flujo de Aire Requerido", "1900 m³/h (1117.8 CFM)"),
        ("Disipación de Potencia (Sección de Potencia)", "4.92 kW")
    ]
    
    for i, row_data in enumerate(heat_data):
        row = heat_table.rows[i]
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        if i == 0:
            set_cell_shading(row.cells[0], "003366")
            set_cell_shading(row.cells[1], "003366")
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    doc.add_heading("Requisitos Ambientales de Operación", level=2)
    
    # Tabla ambiental
    env_table = doc.add_table(rows=6, cols=2)
    env_table.style = 'Table Grid'
    
    env_data = [
        ("Parámetro", "Especificación"),
        ("Temperatura de Operación", "0°C a 40°C (32°F a 104°F)"),
        ("Temperatura Óptima Recomendada", "20°C a 30°C (68°F a 86°F)"),
        ("Humedad Relativa Máxima", "95% (sin condensación)"),
        ("Punto de Rocío Máximo", "37°C (98.6°F)"),
        ("Altitud Máxima sin Derateo", "1000 m sobre nivel del mar")
    ]
    
    for i, row_data in enumerate(env_data):
        row = env_table.rows[i]
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        if i == 0:
            set_cell_shading(row.cells[0], "003366")
            set_cell_shading(row.cells[1], "003366")
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    doc.add_heading("Consideraciones de Ventilación", level=2)
    
    ventilation = """Para asegurar una operación óptima y prolongar la vida útil del equipo:

• Mantener temperatura ambiente entre 20°C y 30°C para máxima confiabilidad
• Operar por encima o debajo de los límites resultará en apagado o rendimiento reducido
• Los límites de temperatura son máximos/mínimos de diseño, NO temperaturas ideales de operación
• La confiabilidad del sistema y expectativa de vida mejoran en el rango óptimo

IMPORTANTE: 
• Los filtros activos requieren intercambio de aire ambiental sin restricciones
• El ambiente debe cumplir con Grado de Contaminación 2
• No debe contener partículas conductivas, cantidades significativas de polvo, o gases corrosivos
• Solo se espera contaminación no conductiva; conductividad temporal por condensación es aceptable"""
    
    doc.add_paragraph(ventilation)
    
    doc.add_heading("Diseño del Sistema de Enfriamiento", level=2)
    
    cooling = """El sistema de aire acondicionado del cuarto eléctrico debe considerar:

• Capacidad para remover la carga térmica de 4.3 kW del equipo
• Flujo de aire sin obstrucciones hacia las rejillas de ventilación del gabinete
• No bloquear la entrada de aire en la parte inferior del gabinete
• Mantener espacio libre alrededor del equipo para circulación de aire
• Si se instalan múltiples unidades en paralelo, sumar las cargas térmicas individuales"""
    
    doc.add_paragraph(cooling)
    
    doc.add_page_break()
    
    # =====================================================
    # SECCIÓN 4: ESPECIFICACIONES DE ESPACIO E INSTALACIÓN
    # =====================================================
    doc.add_heading("4. Especificaciones de Espacio e Instalación", level=1)
    
    doc.add_heading("Dimensiones del Gabinete", level=2)
    
    dimensions_text = """El gabinete PCSP200D5IP54 tiene las siguientes dimensiones:

• Alto: 2100 mm (82.7 pulgadas)
• Ancho: 900 mm (35.4 pulgadas)  
• Profundidad: 600 mm (23.6 pulgadas)
• Peso: 402 kg (886 libras)

El gabinete incluye orejas de levantamiento para izaje con grúa o montacargas."""
    
    doc.add_paragraph(dimensions_text)
    
    doc.add_heading("Requisitos de Espacio para Instalación", level=2)
    
    space_req = """Espacios Mínimos Recomendados:

• Espacio Frontal: Mínimo 900 mm para apertura completa de puertas y acceso de servicio
• Espacio Lateral: Mínimo 100 mm a cada lado para ventilación
• Espacio Trasero: Mínimo 100 mm para circulación de aire
• Espacio Superior: Mínimo 300 mm si se usa entrada de cables superior

Requisitos del Piso:
• Superficie nivelada y capaz de soportar 402 kg de peso
• Puntos de anclaje para fijación al piso (tornillería no incluida)
• Dimensiones de anclaje: 540 mm x 800 mm (consultar dibujo dimensional)"""
    
    doc.add_paragraph(space_req)
    
    doc.add_heading("Requisitos de Instalación Eléctrica", level=2)
    
    doc.add_paragraph("Protección de Sobrecorriente:")
    
    # Tabla de protección
    prot_table = doc.add_table(rows=4, cols=2)
    prot_table.style = 'Table Grid'
    
    prot_data = [
        ("Parámetro", "Valor"),
        ("Ampacidad Mínima del Circuito", "200 A"),
        ("Fusible/Interruptor Mínimo", "250 A"),
        ("Fusible/Interruptor Máximo", "250 A")
    ]
    
    for i, row_data in enumerate(prot_data):
        row = prot_table.rows[i]
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        if i == 0:
            set_cell_shading(row.cells[0], "003366")
            set_cell_shading(row.cells[1], "003366")
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    electrical = """Cableado de Potencia:
• Los cables deben ser dimensionados al 125% de la corriente nominal (250A)
• Usar conductores de cobre con aislamiento para 75°C o 90°C
• Instalar en conduit metálico o cables encapsulados blindados
• El conduit metálico debe conectarse a tierra en el terminal del equipo
• Entrada de cables por placa superior o inferior según configuración

Dispositivo de Corriente Residual (RCD/GFCI):
• Se recomienda dispositivo de mínimo 500 mA debido a alta corriente de fuga en operación normal
• Si se requiere protección menor a 500 mA, se deben abrir los interruptores IT/BP

Puesta a Tierra:
• Conductor de tierra dedicado obligatorio
• Terminal de tierra M8 provisto en el equipo
• Conectar a tierra antes de energizar el equipo"""
    
    doc.add_paragraph(electrical)
    
    doc.add_heading("Transformadores de Corriente (CTs)", level=2)
    
    ct_info = """El sistema requiere transformadores de corriente externos para el monitoreo de la carga:

Requisitos Mínimos de los CTs:
• Relación: Según la corriente máxima de carga
• Secundario: 5A o 1A (configurable)
• Clase de precisión: 0.5 o mejor
• Burden: Compatible con la longitud de cable utilizada

Los CTs deben instalarse en el lado de la fuente de las cargas a compensar."""
    
    doc.add_paragraph(ct_info)
    
    doc.add_page_break()
    
    # =====================================================
    # SECCIÓN 5: CARACTERÍSTICAS Y MEJORAS AL SISTEMA
    # =====================================================
    doc.add_heading("5. Características y Mejoras al Sistema", level=1)
    
    doc.add_heading("Beneficios de la Implementación", level=2)
    
    benefits = """Mejoras en la Calidad de Energía:
• Reducción de THDi a ≤3% (cumplimiento IEEE 519)
• Corrección del factor de potencia (leading o lagging configurable)
• Balanceo de corrientes de fase
• Reducción de fluctuaciones de voltaje relacionadas con procesos

Mejoras en la Confiabilidad del Sistema:
• Eliminación de disparos térmicos en dispositivos de protección
• Reducción del sobrecalentamiento en cables, interruptores y transformadores
• Mayor vida útil de motores y equipos eléctricos
• Reducción de interferencia electromagnética (EMI)

Mejoras Operativas y Económicas:
• Liberación de capacidad del sistema de distribución
• Reducción de pérdidas en líneas por corrientes armónicas
• Eliminación de penalizaciones por bajo factor de potencia
• Reducción de profundidad de caídas de voltaje por inrush de corriente
• Capacidad adicional para conectar más cargas sin ampliar infraestructura"""
    
    doc.add_paragraph(benefits)
    
    doc.add_heading("Características Técnicas del PCSP200D5IP54", level=2)
    
    features = """Funcionalidades Principales:
• Inyección dinámica de corriente en tiempo real
• Compensación de armónicos hasta el orden 51
• Respuesta a fluctuaciones de carga en 2 ciclos (armónicos)
• Respuesta a cambios de factor de potencia en 1/4 de ciclo
• Capacidad de compensación VAR (adelanto o atraso)
• Operación en paralelo de hasta 10 unidades para mayor capacidad

Interfaz y Comunicaciones:
• Pantalla HMI táctil a color integrada
• Comunicación Modbus TCP/IP (Ethernet)
• Comunicación Modbus Serial (RS-485)
• Entradas digitales configurables
• Salidas de contacto seco programables
• Compatible con EcoStruxure Power para monitoreo remoto y análisis"""
    
    doc.add_paragraph(features)
    
    doc.add_heading("Modos de Operación", level=2)
    
    modes = """El AccuSine PCS+ puede operar en los siguientes modos:

1. Cancelación de Armónicos
   Inyecta corriente armónica opuesta para neutralizar armónicos generados por cargas no lineales

2. Corrección del Factor de Potencia
   Inyecta corriente reactiva para ajustar el factor de potencia al valor objetivo configurado

3. Balanceo de Carga
   Equilibra las corrientes entre las tres fases para reducir corrientes de neutro y pérdidas

4. Modo Combinado
   Operación simultánea de las funciones anteriores con priorización configurable según necesidades"""
    
    doc.add_paragraph(modes)
    
    doc.add_heading("Ventajas sobre Soluciones Convencionales", level=2)
    
    advantages = """Comparado con filtros pasivos y otras soluciones tradicionales:

• Sin riesgo de resonancia con el sistema eléctrico
• Adaptación automática a cambios dinámicos en la carga
• Menor espacio físico requerido que bancos de filtros pasivos
• No requiere estudios complejos de armónicos para dimensionamiento
• Fácil expansión mediante unidades adicionales en paralelo
• Mantenimiento reducido (sin capacitores o inductores que reemplazar)
• Respuesta dinámica superior a cualquier filtro pasivo
• Compatible con generadores sin riesgo de autoexcitación
• No se desintona con cambios en el sistema eléctrico"""
    
    doc.add_paragraph(advantages)
    
    doc.add_page_break()
    
    # =====================================================
    # SECCIÓN 6: NORMATIVAS Y CERTIFICACIONES
    # =====================================================
    doc.add_heading("6. Normativas y Certificaciones", level=1)
    
    doc.add_heading("Cumplimiento de Estándares de Diseño", level=2)
    
    standards = """El AccuSine PCS+ PCSP200D5IP54 está diseñado según las siguientes normas:

Normas de Seguridad y Diseño:
• IEC 62477-1 - Requisitos de seguridad para sistemas de conversión de potencia
• IEC 61000-6-2 - Inmunidad EMC para entornos industriales
• IEC 61000-6-4 - Emisiones EMC para entornos industriales
• IEC 61439-1 - Conjuntos de aparamenta de baja tensión
• UL 508 - Equipos de control industrial

Estándares de Calidad de Energía:
• IEEE 519-2014 - Prácticas recomendadas para control de armónicos en sistemas eléctricos
• IEC 61000-3-4 - Límites de emisión de corrientes armónicas
• UK G5/4-1 - Estándar del Reino Unido para conexión de cargas generadoras de armónicos"""
    
    doc.add_paragraph(standards)
    
    doc.add_heading("Certificaciones del Producto", level=2)
    
    # Tabla de certificaciones
    cert_table = doc.add_table(rows=6, cols=2)
    cert_table.style = 'Table Grid'
    
    cert_data = [
        ("Certificación", "Descripción"),
        ("CE", "Conformidad Europea - Cumple directivas de seguridad y EMC"),
        ("UL", "Underwriters Laboratories - Certificación de seguridad para Norteamérica"),
        ("IBC 2015 AC156", "Código sísmico internacional de construcción"),
        ("EAC", "Certificación Euroasiática (Rusia, Kazajistán, Bielorrusia)"),
        ("RCM", "Marca de Cumplimiento Regulatorio - Australia y Nueva Zelanda")
    ]
    
    for i, row_data in enumerate(cert_data):
        row = cert_table.rows[i]
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        if i == 0:
            set_cell_shading(row.cells[0], "003366")
            set_cell_shading(row.cells[1], "003366")
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    doc.add_heading("Resumen de Especificaciones PCSP200D5IP54", level=2)
    
    # Tabla resumen final
    summary_table = doc.add_table(rows=16, cols=2)
    summary_table.style = 'Table Grid'
    
    summary_data = [
        ("Especificación", "Valor"),
        ("Modelo", "PCSP200D5IP54"),
        ("Corriente Nominal", "200 A"),
        ("Voltaje de Red", "380-480 V AC"),
        ("Frecuencia", "50/60 Hz"),
        ("Potencia Reactiva", "166 kVAR @ 480V"),
        ("THDi de Salida", "≤ 3%"),
        ("Grado de Protección", "IP54"),
        ("Dimensiones (A x An x P)", "2100 x 900 x 600 mm"),
        ("Peso", "402 kg"),
        ("Disipación de Calor", "3.9 - 4.3 kW"),
        ("Temperatura de Operación", "0°C a 40°C"),
        ("Montaje", "Piso"),
        ("Entrada de Cables", "Superior o Inferior"),
        ("Certificaciones", "CE, UL, IBC, EAC, RCM"),
        ("Fabricante", "Schneider Electric")
    ]
    
    for i, row_data in enumerate(summary_data):
        row = summary_table.rows[i]
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        if i == 0:
            set_cell_shading(row.cells[0], "003366")
            set_cell_shading(row.cells[1], "003366")
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Nota final
    doc.add_paragraph()
    note = doc.add_paragraph()
    note_run = note.add_run("NOTA: ")
    note_run.bold = True
    note.add_run("Las especificaciones están sujetas a cambios sin previo aviso. Consulte la documentación oficial de Schneider Electric para información actualizada. Este documento es de referencia y no sustituye los manuales de instalación y operación del fabricante.")
    
    # Guardar documento
    doc.save('/workspace/AccuSine_PCS200_Informacion_Tecnica.docx')
    print("Documento creado exitosamente: AccuSine_PCS200_Informacion_Tecnica.docx")

if __name__ == "__main__":
    create_document()
